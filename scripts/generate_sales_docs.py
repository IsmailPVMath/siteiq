#!/usr/bin/env python3
"""Generate PVMath pilot agreement + proforma invoice templates (Word)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
PILOT = DOCS / "PVMath_Pilot_Subscription_Agreement.docx"
INVOICE = DOCS / "PVMath_Proforma_Invoice_Template.docx"
RUNBOOK = DOCS / "PVMath_Manual_Billing_Runbook.md"


def _style(doc: Document) -> None:
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(11)


def build_pilot_agreement() -> Document:
    doc = Document()
    _style(doc)
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("PVMath — Pilot Subscription Agreement\n(Early Access)")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(20, 95, 52)
    doc.add_paragraph("Fill in [brackets]. Bilingual DE/EN for reference — legal version to be confirmed with Steuerberater.")
    doc.add_paragraph()

    sections = [
        ("1. Parties", [
            "Provider: [Mohammed Ismail Pasha / later: PVMath UG (haftungsbeschränkt)], "
            "[Address], Germany, contact@pvmath.com",
            "Customer: [Company legal name], [Address], [Country], [Billing contact email]",
        ]),
        ("2. Service", [
            "Access to the PVMath Solar Site Intelligence Platform (SiteIQ, TopoIQ, YieldIQ) "
            "at siteiq.pvmath.com during Early Access.",
            "Outputs are screening-grade only — not bankable yield, not survey-grade terrain, "
            "not a substitute for LiDAR, PVsyst, or lender sign-off.",
        ]),
        ("3. Plan & limits", [
            "Plan: [ ] Professional (€149/month)  [ ] Developer (€499/month)",
            "Professional: 75 analyses per calendar month, shared across SiteIQ, TopoIQ, and YieldIQ.",
            "Developer: 150 analyses per calendar month, shared team pool (up to 5 seats).",
            "Limits reset on the 1st of each month (UTC). Unused analyses do not roll over.",
        ]),
        ("4. Term & payment", [
            "Pilot term: [3] months from activation date [DD.MM.YYYY].",
            "Payment: bank transfer (SEPA) to [IBAN] within [14] days of this agreement / proforma invoice.",
            "Reference: PVMath-[Customer short name]-[YYYY-MM].",
            "VAT: [19% USt if applicable / reverse charge if EU B2B with valid USt-IdNr.]",
        ]),
        ("5. Activation", [
            "Customer creates user account(s) at siteiq.pvmath.com.",
            "Provider activates plan within 2 business days of cleared payment.",
            "Login email(s) for activation: [list emails]",
        ]),
        ("6. Data & confidentiality", [
            "Customer retains ownership of uploaded boundaries/KMZ and project data.",
            "Provider processes data per privacy policy at pvmath.com/privacy.html.",
            "Optional: separate AVV/DPA on request for enterprise customers.",
        ]),
        ("7. Liability", [
            "Provider liability is limited to fees paid in the 12 months preceding a claim, "
            "except where mandatory law applies.",
            "Customer confirms outputs are for internal go/no-go screening only.",
        ]),
        ("8. Entity change", [
            "If Provider incorporates as PVMath UG (haftungsbeschränkt), this agreement "
            "continues unchanged; invoices will be reissued in the company name upon Handelsregister entry.",
        ]),
        ("9. Cancellation", [
            "Either party may cancel at end of pilot term with [30] days written notice.",
            "No automatic renewal until Stripe or formal subscription is in place.",
        ]),
        ("10. Signatures", [
            "Provider: _________________________  Date: __________  Mohammed Ismail Pasha",
            "Customer: _________________________  Date: __________  Name / Title",
        ]),
    ]
    for title, bullets in sections:
        doc.add_heading(title, level=2)
        for b in bullets:
            doc.add_paragraph(b, style="List Bullet")
    doc.add_page_break()
    doc.add_heading("Deutsch — Kurzfassung (Referenz)", level=1)
    doc.add_paragraph(
        "Pilot-Zugang zu PVMath (SiteIQ, TopoIQ, YieldIQ). Screening-Qualität, keine "
        "Bankfähigkeit. Professional: 75 Analysen/Monat gesamt. Developer: 150/Monat Team-Pool. "
        "Freischaltung nach Zahlungseingang. Bei UG-Gründung nahtlose Fortführung."
    )
    return doc


def build_invoice() -> Document:
    doc = Document()
    _style(doc)
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("PROFORMA INVOICE / PROFORMARECHNUNG")
    r.bold = True
    r.font.size = Pt(18)
    doc.add_paragraph()

    rows = [
        ("Invoice no.", "PV-[YYYY-MM]-[001]"),
        ("Date", "[DD.MM.YYYY]"),
        ("Due date", "[DD.MM.YYYY] (+14 days)"),
        ("", ""),
        ("From / Rechnungssteller", "[Mohammed Ismail Pasha]\n[Street]\n[PLZ City, Germany]\n"
         "contact@pvmath.com\n[Steuernummer when available]\n[USt-IdNr. when available]"),
        ("Bill to / Rechnungsempfänger", "[Company name]\n[Address]\n[USt-IdNr. if EU B2B]"),
        ("", ""),
    ]
    for a, b in rows:
        p = doc.add_paragraph()
        p.add_run(f"{a}\n").bold = bool(a)
        p.add_run(b)

    doc.add_heading("Line items", level=2)
    table = doc.add_table(rows=3, cols=4)
    table.style = "Table Grid"
    h = table.rows[0].cells
    h[0].text = "Description"
    h[1].text = "Period"
    h[2].text = "Net (€)"
    h[3].text = "VAT"
    table.rows[1].cells[0].text = "PVMath Professional — Early Access pilot (75 analyses/month pooled)"
    table.rows[1].cells[1].text = "[Month YYYY-MM]"
    table.rows[1].cells[2].text = "149.00"
    table.rows[1].cells[3].text = "19%"
    table.rows[2].cells[0].text = "PVMath Developer — Early Access pilot (150/month team pool, 5 seats)"
    table.rows[2].cells[1].text = "[Month YYYY-MM]"
    table.rows[2].cells[2].text = "499.00"
    table.rows[2].cells[3].text = "19%"

    doc.add_paragraph()
    doc.add_paragraph("Subtotal (net): €[amount]")
    doc.add_paragraph("VAT 19%: €[amount]  (or: Reverse charge — Art. 196 EU VAT Directive)")
    doc.add_paragraph("Total due: €[amount]")
    doc.add_paragraph()
    doc.add_paragraph("Payment: SEPA bank transfer")
    doc.add_paragraph("IBAN: [DE…]")
    doc.add_paragraph("BIC: […]")
    doc.add_paragraph("Reference: PVMath-[Customer]-[YYYY-MM]")
    doc.add_paragraph()
    doc.add_paragraph(
        "Service activation within 2 business days of payment. Screening-grade outputs only. "
        "Terms: pvmath.com/terms.html"
    )
    return doc


def write_runbook() -> None:
    RUNBOOK.write_text(
        """# PVMath — Manual billing runbook (until Stripe + UG)

## Customer flow

1. User hits upgrade in app → `contact@pvmath.com` (already wired).
2. You reply with **Pilot Subscription Agreement** + **Proforma invoice** (docs folder).
3. Customer signs + pays by bank transfer.
4. You activate plan in Supabase (below).
5. Customer refreshes app — new limits apply immediately.

## Activate plan in Supabase

1. Supabase → Authentication → find user email → copy **User UID**.
2. SQL Editor:

```sql
-- Professional (75 analyses/month pooled across modules)
update profiles set plan = 'professional' where id = '<user-uuid>';

-- Developer (150/month pooled — team shares usage_key; owner uuid = team_id for members)
update profiles set plan = 'developer' where id = '<owner-uuid>';
-- teammates:
update profiles set plan = 'developer', team_id = '<owner-uuid>' where id = '<teammate-uuid>';
```

3. Ask customer to log out/in if plan badge does not update.

## How limits work (Professional & Developer)

- **Professional:** **75 total runs/month** across SiteIQ + TopoIQ + YieldIQ.
- **Developer:** **150 total runs/month** — **entire team shares one pool** (up to 5 seats).
- Example: 60 TopoIQ + 15 SiteIQ = 75 → Professional paywall until next month.
- Enforced in app via `is_over_limit()` — no manual tracking needed.
- Customer sees **X / limit** on **Overview** dashboard.

## Free tier (unchanged)

- 5 runs **per module** per month.

## Verify usage

```sql
select app, count, period from usage_tracking
where usage_key = '<user-uuid>' and period = to_char(now(), 'YYYY-MM');
```

Sum `count` across apps for pooled Professional total.

## Documents

- `PVMath_Pilot_Subscription_Agreement.docx`
- `PVMath_Proforma_Invoice_Template.docx`

Regenerate: `python3 scripts/generate_sales_docs.py`
""",
        encoding="utf-8",
    )


def main() -> None:
    DOCS.mkdir(parents=True, exist_ok=True)
    build_pilot_agreement().save(PILOT)
    build_invoice().save(INVOICE)
    write_runbook()
    print(f"Wrote {PILOT}")
    print(f"Wrote {INVOICE}")
    print(f"Wrote {RUNBOOK}")


if __name__ == "__main__":
    main()
