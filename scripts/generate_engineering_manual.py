#!/usr/bin/env python3
"""
Regenerate PVMath Engineering Reference Manuals (Word + optional PDF).

Usage:
    python3 scripts/generate_engineering_manual.py

Output:
    docs/PVMath_Engineering_Reference_Manual_INTERNAL.docx  — confidential, full IP
    docs/PVMath_Engineering_Reference_Manual_PUBLIC.docx    — customer-safe, redacted
"""
from __future__ import annotations

import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
OUT_INTERNAL = DOCS / "PVMath_Engineering_Reference_Manual_INTERNAL.docx"
OUT_PUBLIC = DOCS / "PVMath_Engineering_Reference_Manual_PUBLIC.docx"

VERSION = "1.1"
MANUAL_DATE = "June 2026 (LayoutIQ update)"

PUBLIC_CALC_REDACTION = (
    "Computed inside PVMath using industry-standard solar and terrain methods on region-routed "
    "public DEM (EEA-10, USGS 3DEP, FABDEM, GLO-30 fallback), PVGIS, and disclosed loss assumptions. "
    "Proprietary weighting, thresholds, and implementation constants are not published in this edition. "
    "Use in-app help and the Knowledge Centre at pvmath.com/guides for interpretation."
)

sys.path.insert(0, str(Path(__file__).parent))
from manual_terms_data import (  # noqa: E402
    ALGORITHM_ENTRIES,
    DATA_SOURCE_ENTRIES,
    EXPORT_ENTRIES,
    GLOSSARY_ENTRIES,
    REPORT_FIELD_ENTRIES,
    SECTIONS,
    TERM_ENTRIES,
)

ENTRY_LABELS = [
    ("definition_plain", "Definition (plain English)"),
    ("definition_formal", "Formal engineering definition"),
    ("why_matters", "Why it matters in utility-scale solar"),
    ("pvmath_calc", "How PVMath calculates or derives it"),
    ("units", "Units"),
    ("typical_ranges", "Typical industry ranges and interpretation"),
    ("example", "Practical example"),
    ("accuracy", "Accuracy considerations and limitations"),
    ("misconceptions", "Common misconceptions"),
    ("references", "References to standards and sources"),
    ("related", "Related PVMath metrics"),
]

PUBLIC_SECTIONS = [s for s in SECTIONS if "Algorithms" not in s["title"]]


def _set_run_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_page_number_footer(section, edition: str):
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"PVMath Engineering Reference Manual ({edition})  ·  v{VERSION}  ·  Page ")
    _set_run_font(run, size=9, color=(100, 100, 100))
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)
    run = p.add_run()
    _set_run_font(run, size=9, color=(100, 100, 100))
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def _style_doc(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for level, size in [(1, 18), (2, 14), (3, 12)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = RGBColor(13, 33, 55)


def _add_cover(doc: Document, *, public: bool):
    for _ in range(5 if public else 6):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = "PVMath Engineering Reference Manual"
    if public:
        title += "\nPublic Edition"
    else:
        title += "\nInternal Edition — CONFIDENTIAL"
    r = t.add_run(title)
    _set_run_font(r, size=26 if public else 24, bold=True, color=(29, 158, 82))
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("SiteIQ  ·  TerrainIQ  ·  LayoutIQ  ·  YieldIQ")
    _set_run_font(r, size=16, color=(26, 46, 26))
    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run(f"Version {VERSION}  ·  {MANUAL_DATE}")
    _set_run_font(r, size=12, color=(80, 80, 80))
    doc.add_paragraph()
    disc = doc.add_paragraph()
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if public:
        txt = (
            "Customer-safe engineering concepts for ground-mount solar screening. "
            "Does not disclose proprietary scoring weights, thresholds, or implementation "
            "constants. For portfolio interpretation use in-app help and pvmath.com/guides."
        )
    else:
        txt = (
            "CONFIDENTIAL — PVMath internal use only. Contains proprietary algorithms, "
            "weights, and thresholds as implemented in the solarscout codebase. "
            "Do not distribute externally."
        )
    r = disc.add_run(txt)
    _set_run_font(r, size=10, color=(180, 0, 0) if not public else (100, 100, 100))
    doc.add_page_break()


def _add_toc(doc: Document, *, public: bool):
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph(
        "Update fields in Microsoft Word (References → Update Table) if automatic "
        "navigation is required; section list below reflects manual structure."
    )
    for sec in (PUBLIC_SECTIONS if public else SECTIONS):
        p = doc.add_paragraph(style="List Number")
        p.add_run(sec["title"]).bold = True
        for child in sec.get("children", []):
            cp = doc.add_paragraph(style="List Bullet")
            cp.paragraph_format.left_indent = Inches(0.35)
            cp.add_run(child)
    doc.add_page_break()


def _prepare_entry(entry: dict, *, public: bool) -> dict:
    if not public:
        return entry
    out = dict(entry)
    if out.get("pvmath_calc", "").strip():
        out["pvmath_calc"] = PUBLIC_CALC_REDACTION
    return out


def _write_entry(doc: Document, entry: dict, level: int = 2):
    doc.add_heading(entry["title"], level=level)
    for key, label in ENTRY_LABELS:
        val = entry.get(key, "").strip()
        if not val:
            continue
        p = doc.add_paragraph()
        r = p.add_run(f"{label}: ")
        r.bold = True
        p.add_run(val)
    doc.add_paragraph()


def _write_section(doc: Document, title: str, intro: str, entries: list[dict], *, public: bool):
    doc.add_heading(title, level=1)
    if intro:
        doc.add_paragraph(intro)
    for e in entries:
        _write_entry(doc, _prepare_entry(e, public=public))
    doc.add_page_break()


def build_document(*, public: bool = False) -> Document:
    doc = Document()
    _style_doc(doc)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    edition = "Public" if public else "Internal"
    _add_page_number_footer(section, edition)

    _add_cover(doc, public=public)
    _add_toc(doc, public=public)

    _write_section(
        doc, "Terrain Analysis (TerrainIQ)", SECTIONS[0]["intro"],
        [t for t in TERM_ENTRIES if t["section"] == "terrain"], public=public,
    )
    _write_section(
        doc, "Solar Resource & Site Screening (SiteIQ)", SECTIONS[1]["intro"],
        [t for t in TERM_ENTRIES if t["section"] == "solar"], public=public,
    )
    _write_section(
        doc, "Yield & Energy (YieldIQ)", SECTIONS[2]["intro"],
        [t for t in TERM_ENTRIES if t["section"] == "yield"], public=public,
    )
    _write_section(
        doc, "Layout, Grading & Capacity", SECTIONS[3]["intro"],
        [t for t in TERM_ENTRIES if t["section"] == "capacity"], public=public,
    )
    _write_section(
        doc, "GIS, CAD & Export Formats", SECTIONS[4]["intro"],
        EXPORT_ENTRIES, public=public,
    )
    _write_section(
        doc, "Data Sources & APIs", SECTIONS[5]["intro"],
        DATA_SOURCE_ENTRIES, public=public,
    )
    _write_section(
        doc, "Report Metrics & PDF Fields", SECTIONS[6]["intro"],
        REPORT_FIELD_ENTRIES, public=public,
    )
    if not public:
        _write_section(
            doc, "Algorithms & Constants", SECTIONS[7]["intro"],
            ALGORITHM_ENTRIES, public=public,
        )

    doc.add_heading("Glossary (Alphabetical)", level=1)
    doc.add_paragraph(
        "Concise alphabetical index of key terms. Full entries appear in the sections above."
    )
    for g in GLOSSARY_ENTRIES:
        p = doc.add_paragraph()
        r = p.add_run(f"{g['term']} — ")
        r.bold = True
        p.add_run(g["brief"])

    doc.add_page_break()
    doc.add_heading("Index", level=1)
    index_sources = TERM_ENTRIES + EXPORT_ENTRIES + DATA_SOURCE_ENTRIES + REPORT_FIELD_ENTRIES
    if not public:
        index_sources += ALGORITHM_ENTRIES
    index_terms = sorted(
        {t["title"].split("(")[0].strip() for t in TERM_ENTRIES}
        | {e["title"].split("(")[0].strip() for e in index_sources}
        | {g["term"] for g in GLOSSARY_ENTRIES},
        key=str.lower,
    )
    cols = 3
    per_col = (len(index_terms) + cols - 1) // cols
    table = doc.add_table(rows=per_col, cols=cols)
    for i, term in enumerate(index_terms):
        row, col = divmod(i, cols)
        if row < per_col:
            table.rows[row].cells[col].text = term

    return doc


def try_pdf(docx_path: Path, pdf_path: Path) -> bool:
    converters = [
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(pdf_path.parent), str(docx_path)],
        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(pdf_path.parent), str(docx_path)],
    ]
    for cmd in converters:
        try:
            subprocess.run(cmd, check=True, capture_output=True, timeout=120)
            generated = pdf_path.parent / f"{docx_path.stem}.pdf"
            if generated.exists() and generated != pdf_path:
                generated.rename(pdf_path)
            if pdf_path.exists():
                return True
        except (FileNotFoundError, subprocess.CalledProcessError, subprocess.TimeoutExpired):
            continue
    try:
        from docx2pdf import convert  # type: ignore
        convert(str(docx_path), str(pdf_path))
        return pdf_path.exists()
    except Exception:
        return False


def main():
    DOCS.mkdir(parents=True, exist_ok=True)
    for path, public in ((OUT_INTERNAL, False), (OUT_PUBLIC, True)):
        doc = build_document(public=public)
        doc.save(str(path))
        label = "PUBLIC" if public else "INTERNAL"
        print(f"Wrote {path} ({label})")

    term_count = (
        len(TERM_ENTRIES)
        + len(EXPORT_ENTRIES)
        + len(DATA_SOURCE_ENTRIES)
        + len(REPORT_FIELD_ENTRIES)
        + len(ALGORITHM_ENTRIES)
    )
    print(f"Distinct entries: {term_count} (+ {len(GLOSSARY_ENTRIES)} glossary lines)")
    print(f"Public edition: algorithms section omitted; pvmath_calc fields redacted")


if __name__ == "__main__":
    main()
