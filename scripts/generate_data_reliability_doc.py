#!/usr/bin/env python3
"""Generate PVMath data sources & availability brief (Word) — sales / pilot use."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "PVMath_Data_Sources_and_Availability.docx"


def _style(doc: Document) -> None:
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(11)


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build() -> Document:
    doc = Document()
    _style(doc)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("PVMath — Data Sources & Availability\n(Sales & pilot briefing)")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(20, 95, 52)

    doc.add_paragraph(
        "Internal / customer-facing attachment for EPC and developer conversations. "
        "Explains what PVMath depends on, what paying customers receive, and honest limits "
        "when third-party APIs are disrupted. Not a contractual SLA — pair with the Pilot "
        "Subscription Agreement for paid pilots."
    )
    doc.add_paragraph("Last updated: June 2026 · contact@pvmath.com · pvmath.com")

    doc.add_heading("1. Executive summary", level=1)
    doc.add_paragraph(
        "PVMath (SiteIQ, TopoIQ, YieldIQ) is a screening and workflow layer on top of "
        "public scientific datasets that the solar industry already uses — principally "
        "PVGIS (EU JRC) for irradiance and yield, Copernicus GLO-30 (via AWS Terrarium "
        "tiles) and SRTM/EU-DEM (OpenTopoData) for terrain. PVMath does not host or "
        "mirror those datasets today."
    )
    doc.add_paragraph(
        "Paying customers buy integrated analysis, saved projects, PDF/CAD outputs, pooled "
        "usage limits, and support — not guaranteed uptime of EU or AWS infrastructure. "
        "Outputs remain screening-grade (pre-survey, pre-bankable), as stated on reports "
        "and in the pilot agreement."
    )

    doc.add_heading("2. External data sources", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Source"
    hdr[1].text = "Used for"
    hdr[2].text = "Operator"
    hdr[3].text = "PVMath status today"
    rows = [
        (
            "PVGIS (JRC / EU Commission)",
            "Solar resource, YieldIQ, SiteIQ capacity context",
            "Public EU research API",
            "Primary; automatic fallback to PVGIS-ERA5 when primary DB unavailable",
        ),
        (
            "Copernicus GLO-30 via AWS Terrarium",
            "TopoIQ DEM, slope maps, CAD exports",
            "AWS open elevation tile CDN",
            "Single live source; partial tile failure tolerated; no mirror archive",
        ),
        (
            "OpenTopoData (EU-DEM 25 m / SRTM 30 m)",
            "SiteIQ quick terrain (pin mode)",
            "Public API",
            "Single source; dataset chosen by coordinates",
        ),
        (
            "Nominatim / OpenStreetMap",
            "Place search, location labels on reports",
            "Public geocoding",
            "Non-critical; analysis works from coordinates alone",
        ),
        (
            "Google map tiles",
            "Map display in Project Setup",
            "Google",
            "Display only — not used in calculations",
        ),
    ]
    for src, use, op, status in rows:
        row = table.add_row().cells
        row[0].text = src
        row[1].text = use
        row[2].text = op
        row[3].text = status

    doc.add_paragraph()
    doc.add_heading("3. What resilience exists today", level=1)
    _bullets(doc, [
        "Solar / yield: PVGIS-ERA5 fallback implemented in production code when the "
        "preferred radiation database fails for a site.",
        "TopoIQ: individual DEM tiles that fail or time out are skipped (nodata) rather "
        "than crashing the whole run; georeferencing is preserved.",
        "TopoIQ → SiteIQ cache: after a successful TopoIQ run on a saved project boundary, "
        "SiteIQ can reuse confirmed GLO-30 terrain from cache — repeat screening does not "
        "always need live tiles.",
        "PDF disclaimers: all modules label outputs as screening-grade, not survey or "
        "bankable yield.",
        "Platform: app hosted on Railway; auth and projects on Supabase — separate from "
        "scientific data APIs.",
    ])

    doc.add_heading("4. What we do not guarantee today", level=1)
    _bullets(doc, [
        "Third-party API uptime (PVGIS, AWS Terrarium, OpenTopoData, Nominatim).",
        "Continuous availability of TopoIQ during a multi-day AWS elevation tile outage.",
        "That cached terrain replaces the need for a fresh TopoIQ run when boundaries change.",
        "Enterprise SLA on data providers — only the PVMath application layer can be "
        "addressed in future enterprise contracts.",
    ])

    doc.add_heading("5. Scenario: AWS elevation tiles unavailable ~48 hours", level=1)
    doc.add_paragraph(
        "Honest customer answer:"
    )
    doc.add_paragraph(
        "“PVMath stays online, but TopoIQ terrain extraction and any new analysis that "
        "requires live Copernicus/AWS DEM tiles will fail or be incomplete until the tile "
        "service recovers. SiteIQ solar screening and YieldIQ may still work if PVGIS is "
        "available. PDFs and CAD files already downloaded remain with the customer. We do "
        "not delete completed work. This is the same dependency any team faces when calling "
        "public APIs directly — PVMath adds workflow and reports on top, not a private DEM "
        "archive.”"
    )

    doc.add_heading("6. What paying customers are buying", level=1)
    table2 = doc.add_table(rows=1, cols=2)
    table2.style = "Table Grid"
    table2.rows[0].cells[0].text = "Self-serve / free tier"
    table2.rows[0].cells[1].text = "Professional / Developer (paid pilot)"
    paid_rows = [
        ("Same underlying open data", "Pooled analyses (75 / 300 per month), team seats"),
        ("Manual PVGIS + GIS workflow", "Integrated SiteIQ + TopoIQ + YieldIQ pipeline"),
        ("No support channel", "contact@pvmath.com, pilot agreement, manual activation"),
        ("—", "Saved projects, PDF reports, CAD/LandXML exports"),
    ]
    for a, b in paid_rows:
        row = table2.add_row().cells
        row[0].text = a
        row[1].text = b

    doc.add_paragraph()
    doc.add_heading("7. Roadmap (direction — not live promises)", level=1)
    _bullets(doc, [
        "In-app status when PVGIS / tiles / OpenTopoData fail; retries with backoff.",
        "Optional OpenTopoData fallback for TopoIQ when Terrarium is down (lower resolution, "
        "analysis continues).",
        "Enterprise: SLA on PVMath platform availability and support response — not on JRC/AWS.",
        "Longer term: selective caching or licensed mirrors only if customer demand justifies cost.",
    ])

    doc.add_heading("8. Suggested one-paragraph reply (email / call)", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "“PVMath sits on the same open scientific stack the industry already trusts — PVGIS "
        "for irradiance, Copernicus/SRTM for terrain — and adds workflow, reports, and "
        "project context on top. We are transparent that we do not control those APIs: if "
        "PVGIS or AWS elevation tiles have an outage, affected modules pause until they "
        "recover, the same as if you called those APIs yourself. What you pay for is "
        "integrated screening across SiteIQ, TopoIQ, and YieldIQ, saved projects, PDF/CAD "
        "outputs, and support — plus we already fall back solar to ERA5 and cache TopoIQ "
        "terrain for reuse on the same boundary. Enterprise can include platform SLA and a "
        "defined response when third-party data is disrupted; we are not claiming to "
        "replace LiDAR survey or a private weather feed.”"
    ).italic = True

    doc.add_heading("9. Knowledge Centre — optional public page", level=1)
    doc.add_paragraph(
        "Recommendation: keep this Word document for sales, pilots, and email attachments. "
        "Add a shorter public page later (e.g. guides/data-sources.html) if prospects ask "
        "repeatedly — trim sections 4–5 to customer-friendly language; do not publish "
        "internal roadmap details. A public page builds trust; this document carries the "
        "full talking points."
    )

    doc.add_paragraph()
    doc.add_paragraph(
        "Regenerate: python3 scripts/generate_data_reliability_doc.py"
    ).runs[0].font.size = Pt(9)

    return doc


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    build().save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
