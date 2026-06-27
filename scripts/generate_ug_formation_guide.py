#!/usr/bin/env python3
"""Generate PVMath UG formation guide (Word)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "PVMath_UG_Formation_Guide.docx"


def _style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for level, size in [(1, 18), (2, 14), (3, 12)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = RGBColor(20, 95, 52)


def _p(doc: Document, text: str, bold: bool = False) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = bold


def _bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def build() -> Document:
    doc = Document()
    _style_doc(doc)

    # Cover
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("PVMath UG Formation Guide")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(20, 95, 52)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run("From sole operator to PVMath UG (haftungsbeschränkt)\nJune 2026")
    s.font.size = Pt(12)
    s.font.color.rgb = RGBColor(90, 90, 90)
    doc.add_page_break()

    _p(
        doc,
        "Disclaimer: Practical guidance based on the current PVMath setup — not legal or tax "
        "advice. For UG formation, employment contract review, and VAT, use a Steuerberater and "
        "ideally a Fachanwalt für Gesellschaftsrecht in Bavaria.",
        bold=True,
    )
    doc.add_paragraph()

    doc.add_heading("Where you are today", level=1)
    rows = [
        ("Legal operator", "Mohammed Ismail Pasha personally — Impressum/Terms, Regensburg"),
        ("Product", "SiteIQ, TerrainIQ, YieldIQ live at siteiq.pvmath.com / topoiq.pvmath.com"),
        ("Website", "pvmath.com — Professional €149/mo, Developer €499/mo, VAT noted"),
        ("Payments", "Stripe not live yet — good timing to open under the UG"),
        ("Job", "Full-time solar engineer at Ideematec GmbH — side project"),
        ("Domains", "pvmath.com, pvmath.de, pvmath.eu"),
    ]
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Current state"
    for i, (a, b) in enumerate(rows, start=1):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
    doc.add_paragraph()
    _p(
        doc,
        "A UG (haftungsbeschränkt) fits: low start capital (from €1), limited liability, "
        "credible for B2B EPCs, path to GmbH later.",
    )

    doc.add_heading("Phase 0 — Before you spend money (1–2 weeks)", level=1)

    doc.add_heading("Step 1: Check your employment contract", level=2)
    _p(doc, "This comes first.")
    _bullet(doc, "Read Nebentätigkeitsklausel, Wettbewerbsverbot, IP clause")
    _bullet(doc, "PVMath (solar screening SaaS) may overlap with EPC work — disclose Nebentätigkeit in writing")
    _bullet(doc, "Get written approval before registering the UG if required")
    _bullet(doc, "Keep PVMath outside work hours and on personal equipment until approved")
    _p(doc, "Without this, a UG can create employment risk.")

    doc.add_heading("Step 2: Decide company details", level=2)
    _p(doc, "Suggested legal name: PVMath UG (haftungsbeschränkt)")
    _bullet(doc, "Sitz (registered office): Regensburg")
    _bullet(doc, "Geschäftsführer: Mohammed Ismail Pasha (sole founder)")
    _bullet(doc, "Stammkapital: €1,000–€2,500 typical for SaaS; minimum legal is €1")
    _bullet(
        doc,
        "Unternehmensgegenstand: e.g. Entwicklung und Vertrieb von Software für die technische "
        "Vorprüfung und Standortanalyse von Photovoltaik-Freiflächenanlagen (Steuerberater refines)",
    )
    _bullet(
        doc,
        "UG rule: 25% of annual profit retained until Stammkapital reaches €25,000 (optional GmbH conversion later)",
    )

    doc.add_heading("Step 3: Book professionals", level=2)
    _bullet(doc, "Steuerberater (Regensburg or remote) — formation, VAT, bookkeeping")
    _bullet(doc, "Notar — Ein-personen-UG via Musterprotokoll (~€300–€500 for solo founder)")
    _bullet(doc, "Optional: lawyer for SaaS AGB beyond current English terms")
    _p(doc, "Budget ballpark: €800–€2,000 setup + ~€100–€250/month accounting once revenue starts.")

    doc.add_heading("Phase 1 — Form the UG (2–4 weeks)", level=1)

    doc.add_heading("Step 4: Notary — Gesellschaftsgründung", level=2)
    _p(doc, "Bring: ID, Meldebescheinigung, company name (check IHK/Unternehmensregister), "
         "registered address, Stammkapital amount.")
    _bullet(doc, "Notarise Gesellschaftsvertrag (Musterprotokoll for standard solo UG)")
    _bullet(doc, "You become Geschäftsführer")
    _bullet(doc, "Handelsregister application prepared")

    doc.add_heading("Step 5: Pay Stammkapital", level=2)
    _bullet(doc, "Open Geschäftskonto (see Phase 3)")
    _bullet(doc, "Transfer Stammkapital with reference Stammkapitaleinlage PVMath UG")
    _bullet(doc, "Bank issues Einzahlungsbestätigung → back to notary/register")

    doc.add_heading("Step 6: Handelsregister entry", level=2)
    _bullet(doc, "Notary submits to Amtsgericht Regensburg")
    _bullet(doc, "Wait for HRB number — often 1–3 weeks")
    _bullet(doc, "May operate as PVMath UG (haftungsbeschränkt) i.G. before entry; B2B prefers full HRB")

    doc.add_heading("Phase 2 — Gewerbe & tax (same month)", level=1)

    doc.add_heading("Step 7: Gewerbeanmeldung", level=2)
    _p(doc, "At Stadt Regensburg Gewerbeamt (or online): Software / IT-Dienstleistungen. Fee ~€20–€60.")

    doc.add_heading("Step 8: Finanzamt — Fragebogen zur steuerlichen Erfassung", level=2)
    tax = doc.add_table(rows=5, cols=2)
    tax.style = "Table Grid"
    tax.rows[0].cells[0].text = "Topic"
    tax.rows[0].cells[1].text = "Recommendation for PVMath"
    tax.rows[1].cells[0].text = "Umsatzsteuer"
    tax.rows[1].cells[1].text = "Regular VAT (19%) — site already mentions VAT for EU; not Kleinunternehmer §19"
    tax.rows[2].cells[0].text = "USt-IdNr."
    tax.rows[2].cells[1].text = "Request DE… VAT ID for B2B EU invoices"
    tax.rows[3].cells[0].text = "Gewerbesteuer"
    tax.rows[3].cells[1].text = "Applies once profit exceeds Freibetrag (~€24,500)"
    tax.rows[4].cells[0].text = "Geschäftsführer salary"
    tax.rows[4].cells[1].text = "Steuerberater models Gehalt vs. Gewinnausschüttung"

    doc.add_heading("Step 9: IHK membership", level=2)
    _p(doc, "Software UG in Bavaria → IHK membership mandatory (~€50–€150/year initially).")

    doc.add_heading("Phase 3 — Business infrastructure", level=1)

    doc.add_heading("Step 10: Business bank account", level=2)
    _p(doc, "Options: Kontist, Finom, N26 Business, Deutsche Bank, etc.")
    _p(doc, "Use for: Stripe, Railway, Supabase, Namecheap, Brevo, all SaaS costs.")

    doc.add_heading("Step 11: Accounting setup", level=2)
    _bullet(doc, "Lexoffice, sevDesk, or Steuerberater with DATEV")
    _bullet(doc, "Separate business vs. personal expenses from day one")

    doc.add_heading("Step 12: Insurance", level=2)
    _bullet(doc, "Betriebshaftpflicht with IT-/Software-Haftung (recommended before first paid customer)")
    _bullet(doc, "Optional: Cyber-Versicherung for paying EU enterprise clients")

    doc.add_heading("Phase 4 — Move PVMath assets into the UG", level=1)

    doc.add_heading("Step 13: IP & asset transfer", level=2)
    assets = doc.add_table(rows=6, cols=2)
    assets.style = "Table Grid"
    assets.rows[0].cells[0].text = "Asset"
    assets.rows[0].cells[1].text = "Action"
    asset_rows = [
        ("GitHub repo (IsmailPVMath/siteiq)", "Transfer to UG org or exclusive license to UG"),
        ("Domains (pvmath.com/.de/.eu)", "Transfer registrant to PVMath UG at Namecheap"),
        ("Code, brand, logo", "IP-Übertragungsvertrag from founder → UG"),
        ("Supabase / Railway / Brevo", "Update billing entity to UG"),
        ("Stripe (when live)", "Register under UG name + HRB"),
    ]
    for i, (a, b) in enumerate(asset_rows, start=1):
        assets.rows[i].cells[0].text = a
        assets.rows[i].cells[1].text = b

    doc.add_heading("Step 14: Contract stack for B2B SaaS", level=2)
    _bullet(doc, "AGB / Terms — contracting party = UG, liability limits for screening outputs")
    _bullet(doc, "Datenschutzerklärung — controller = UG, Supabase processor")
    _bullet(doc, "Impressum — full TMG block with HRB and USt-IdNr.")
    _bullet(doc, "AVV/DPA for enterprise customers")
    _p(doc, "Keep existing screening disclaimers (Knowledge Centre, non-bankability terms).")

    doc.add_heading("Phase 5 — Update pvmath.com", level=1)

    doc.add_heading("Step 15: Impressum (after HRB)", level=2)
    impressum = doc.add_paragraph()
    impressum.add_run(
        "PVMath UG (haftungsbeschränkt)\n"
        "[Straße, PLZ Regensburg]\n"
        "Deutschland\n\n"
        "Geschäftsführer: Mohammed Ismail Pasha\n"
        "Handelsregister: Amtsgericht Regensburg, HRB [number]\n"
        "USt-IdNr.: DE[number]"
    )
    _p(doc, "Update: impressum.html, privacy.html, terms.html, index.html footer.")

    doc.add_heading("Step 16: Stripe & invoicing", level=2)
    _bullet(doc, "Stripe account → PVMath UG, German entity")
    _bullet(doc, "Products: Professional €149, Developer €499 (net + 19% VAT DE B2C; reverse charge EU B2B)")
    _bullet(doc, "Invoice template: UG details, Leistungszeitraum, net/gross, USt")
    _bullet(doc, "Wire STRIPE_LINK in pvmath_auth.py when ready")

    doc.add_heading("Phase 6 — Operating rhythm (ongoing)", level=1)

    doc.add_heading("Monthly", level=2)
    _bullet(doc, "Bookkeeping / Steuerberater upload")
    _bullet(doc, "VAT prepayment if applicable")
    _bullet(doc, "Review usage limits, support, Railway costs")

    doc.add_heading("Annually", level=2)
    _bullet(doc, "Jahresabschluss + Offenlegung (Steuerberater advises on Bundesanzeiger)")
    _bullet(doc, "Renew domains, insurance")
    _bullet(doc, "Update Engineering Manual / public guides if product changes")

    doc.add_heading("UG-specific", level=2)
    _bullet(doc, "Reserve 25% of profit until €25k Stammkapital (mandatory)")
    _bullet(doc, "Consider GmbH conversion when revenue justifies €25k capital")

    doc.add_heading("UG now vs cheaper alternatives", level=1)
    _p(
        doc,
        "With PVMath today (live product, Stripe not yet live, early access, side project), "
        "a UG is not urgent. It becomes worth it when you start invoicing B2B customers or "
        "want limited liability on paid subscriptions.",
    )
    alt = doc.add_table(rows=5, cols=5)
    alt.style = "Table Grid"
    alt.rows[0].cells[0].text = "Path"
    alt.rows[0].cells[1].text = "Cost / speed"
    alt.rows[0].cells[2].text = "Liability"
    alt.rows[0].cells[3].text = "B2B credibility"
    alt.rows[0].cells[4].text = "When it fits"
    alt_rows = [
        (
            "Stay private (no Gewerbe)",
            "Free",
            "Personal",
            "Low for paid SaaS",
            "Free tier only, no invoices",
        ),
        (
            "Einzelunternehmen + Gewerbe",
            "~€20–60, few days",
            "Personal (unlimited)",
            "OK for small invoices",
            "First paid tests, low volume",
        ),
        (
            "UG (haftungsbeschränkt)",
            "~€800–2,000, 6–10 weeks",
            "Limited",
            "Good for EPCs / €149–499 subs",
            "Ready to sell seriously",
        ),
        (
            "GmbH",
            "€25k+ capital",
            "Limited",
            "Best for large contracts",
            "Later, if revenue justifies it",
        ),
    ]
    for i, row in enumerate(alt_rows, start=1):
        for j, val in enumerate(row):
            alt.rows[i].cells[j].text = val
    doc.add_paragraph()
    doc.add_heading("Recommended sequence (2026)", level=2)
    _bullet(doc, "Wait for Ideematec Nebentätigkeit approval (in process)")
    _bullet(doc, "Do not rush UG until Stripe live, first paying customer, or EPC asks for company invoice")
    _bullet(
        doc,
        "Cheaper bridge: Gewerbeanmeldung as Einzelunternehmer when first payment is close — "
        "convert to UG when revenue or liability justifies ~€1–2k setup",
    )
    _p(
        doc,
        "Bottom line: UG is worth it when you sell, not because the app is live. "
        "If it stays free-tier only for 12+ months, UG early is mostly admin cost.",
    )

    doc.add_heading("Future: Products vs Services vs Survey", level=1)
    _p(
        doc,
        "Long-term vision: PVMath SaaS + detailed layout / project design services worldwide "
        "+ (after drone licence A2/C1 and capital) LiDAR survey services on the website. "
        "These are three business lines under one brand — not one operational model.",
    )

    doc.add_heading("Three lines — different risk and regulation", level=2)
    lines = doc.add_table(rows=4, cols=3)
    lines.style = "Table Grid"
    lines.rows[0].cells[0].text = "Line"
    lines.rows[0].cells[1].text = "What it is"
    lines.rows[0].cells[2].text = "Risk / regulation"
    line_rows = [
        ("PVMath SaaS", "Screening software (SiteIQ, TerrainIQ, YieldIQ)", "Product liability, AGB, GDPR"),
        (
            "Engineering services",
            "Layout, project design, worldwide consulting",
            "Professional engineering liability, project SOW per job",
        ),
        (
            "LiDAR / drone survey",
            "Field data collection for solar sites",
            "EU drone law, aviation insurance, equipment capex, survey-grade claims",
        ),
    ]
    for i, row in enumerate(line_rows, start=1):
        for j, val in enumerate(row):
            lines.rows[i].cells[j].text = val

    doc.add_heading("Does this change company formation?", level=2)
    _bullet(
        doc,
        "No new company type required — expand Unternehmensgegenstand at formation or later "
        "(Satzungsänderung if needed)",
    )
    _bullet(
        doc,
        "Example broad Gegenstand (Steuerberater refines): Softwareentwicklung; technische Beratung "
        "und Planung für Photovoltaik-Freiflächenanlagen; geoinformationsbezogene Dienstleistungen",
    )
    _bullet(doc, "Drone A2/C1 is operator certification — not company formation")
    _bullet(
        doc,
        "LiDAR needs separate capital (often €50k–200k+ useful kit), insurance, and clear "
        "deliverable scope — official Vermessung in Germany is regulated; solar screening terrain "
        "may be positioned differently (lawyer + Steuerberater define wording)",
    )
    _bullet(
        doc,
        "Same UG is usually fine — but separate contracts and separate insurance per line are mandatory",
    )

    doc.add_heading("Phased approach", level=2)
    doc.add_heading("Phase A — next 12 months (while at Ideematec)", level=3)
    _bullet(doc, "PVMath = SaaS only on website")
    _bullet(doc, "Optional consulting as contact@ only — no full Services page until entity + insurance ready")
    _bullet(doc, "Form UG when SaaS invoices start (or Einzelunternehmer first if Steuerberater agrees)")

    doc.add_heading("Phase B — leave Ideematec, add Services section", level=3)
    _bullet(doc, "Same PVMath UG (or GmbH later) with two contract types: SaaS subscription vs engineering SOW")
    _bullet(doc, "Website: Products | Services — same brand, different terms")
    _bullet(doc, "Professional indemnity / Berufshaftpflicht for engineering work")
    _p(doc, "Services (layout/design) will likely pay faster than subscriptions early on.")

    doc.add_heading("Phase C — drone + LiDAR (after licence + capital)", level=3)
    _bullet(doc, "Option 1: Same UG, division PVMath Survey — separate insurance and SOW")
    _bullet(doc, "Option 2: Second entity (survey-only UG) if lawyer wants liability ring-fenced")
    _bullet(doc, "Do not list LiDAR on website before licence, insurance, entity, and scope are ready")
    _p(doc, "Many founders use one UG first and split only when LiDAR revenue and risk are real.")

    doc.add_heading("Entity structure options", level=2)
    _bullet(
        doc,
        "Option A — One UG, broad purpose: simplest admin; mixed liability — mitigate with insurance + SOWs",
    )
    _bullet(
        doc,
        "Option B — PVMath (software) + separate engineering brand: cleaner insurance, more admin",
    )
    _bullet(
        doc,
        "Option C — One UG, separate contract templates per line (recommended for product-led founders): "
        "SaaS T&Cs, engineering project contract, survey SOW",
    )

    doc.add_heading("What NOT to do", level=2)
    _bullet(doc, "Quit Ideematec on SaaS MRR alone — unlikely to replace salary quickly")
    _bullet(doc, "Advertise survey-grade or bankable LiDAR before you can deliver and insure it")
    _bullet(doc, "Mix screening-grade SaaS language with field-survey promises on the same page without disclaimers")
    _bullet(doc, "Over-engineer two companies at UG formation — expand when LiDAR is real")

    doc.add_heading("Decision tree", level=2)
    _p(doc, "Ideematec approval received?")
    _bullet(doc, "Yes → build product + first users; defer UG until first invoice or Stripe")
    _bullet(doc, "First paying customer? → UG (or Einzelunternehmer → UG within 6 months)")
    _p(doc, "Leave Ideematec?")
    _bullet(doc, "When services pipeline OR SaaS MRR covers 6–12 months runway")
    _p(doc, "Add LiDAR?")
    _bullet(doc, "When A2/C1 + insurance + equipment fund + deliverable scope defined")
    _bullet(doc, "Same UG with expanded Gegenstand usually fine; separate insurance mandatory")

    doc.add_heading("Suggested timeline", level=1)
    _bullet(doc, "Week 1–2: Employment clearance + Steuerberater + Notar booking")
    _bullet(doc, "Week 3–4: Notar + Stammkapital + Gewerbe + Finanzamt")
    _bullet(doc, "Week 5–7: Handelsregister HRB")
    _bullet(doc, "Week 8–10: Bank + asset transfer + website + Stripe")
    _p(doc, "Realistic total: 6–10 weeks from go to first invoice as UG.")

    doc.add_heading("What NOT to change", level=1)
    _bullet(doc, "App code / Railway / modules — unaffected by UG formation")
    _bullet(doc, "Product claims — keep screening-grade honesty")
    _bullet(doc, "Pricing — can stay; invoice as UG with VAT")

    doc.add_heading("Immediate next 3 actions", level=1)
    _bullet(doc, "Complete Ideematec Nebentätigkeit approval (in process — expected ~next week)")
    _bullet(doc, "After approval: email 2 Steuerberäter — Einzelunternehmer vs UG for SaaS + future services")
    _bullet(doc, "Defer UG until Stripe live or first B2B invoice; book IHK name check when ready")

    doc.add_paragraph()
    _p(
        doc,
        "Generated for PVMath / Mohammed Ismail Pasha · pvmath.com · "
        "For personal planning only — confirm all steps with Steuerberater and Notar.",
    )

    return doc


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    build().save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
