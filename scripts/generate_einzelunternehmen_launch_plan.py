#!/usr/bin/env python3
"""Generate PVMath Einzelunternehmen + Intersolar launch plan (Word, EN/DE)."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "PVMath_Einzelunternehmen_Launch_Plan.docx"

GREEN = RGBColor(20, 95, 52)
MUTED = RGBColor(90, 90, 90)


def _style(doc: Document) -> None:
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(11)
    for level, size in [(1, 18), (2, 14), (3, 12)]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = "Calibri"
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = GREEN


def _p(doc: Document, text: str, bold: bool = False, center: bool = False) -> None:
    para = doc.add_paragraph()
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = bold


def _bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def _bilingual_block(doc: Document, en: str, de: str) -> None:
    _p(doc, en)
    p = doc.add_paragraph()
    r = p.add_run(de)
    r.italic = True
    r.font.color.rgb = MUTED


def _table(doc: Document, headers: tuple[str, ...], rows: list[tuple[str, ...]]) -> None:
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = val
    doc.add_paragraph()


def build() -> Document:
    doc = Document()
    _style(doc)

    # ── Cover ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("PVMath — Einzelunternehmen Launch Plan\nIntersolar Week · June 2026")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = GREEN
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run(
        "Mohammed Ismail Pasha · Regensburg · contact@pvmath.com\n"
        "Target: registered sole proprietorship by Saturday evening · "
        "Stripe live · marketing from Tuesday at Intersolar"
    )
    s.font.size = Pt(11)
    s.font.color.rgb = MUTED
    doc.add_page_break()

    _p(
        doc,
        "Disclaimer: Operational checklist for PVMath — not legal or tax advice. "
        "Confirm Gewerbe wording, VAT status (§19 Kleinunternehmer vs regular USt), "
        "Nebentätigkeit, and Stripe invoicing with your Steuerberater before filing.",
        bold=True,
    )
    doc.add_paragraph()

    # ── Mission timeline ─────────────────────────────────────────────────────
    doc.add_heading("Mission timeline / Zeitplan", level=1)
    _table(
        doc,
        ("When / Wann", "Goal / Ziel", "Owner"),
        [
            (
                "Sat 20 Jun (evening)",
                "Gewerbe angemeldet · Fragebogen zur steuerlichen Erfassung submitted · "
                "Steuerberater briefed",
                "You + Steuerberater",
            ),
            (
                "Sun 21 Jun",
                "Stripe account (Einzelunternehmer) · Payment Links · website Subscribe button · "
                "optional webhook scaffold",
                "You + dev",
            ),
            (
                "Mon 22 Jun",
                "Test payment end-to-end · Impressum update · pilot invoice template filled · "
                "LinkedIn / print assets ready",
                "You",
            ),
            (
                "Tue 23 Jun (Intersolar)",
                "Live demos at fair · QR to siteiq.pvmath.com · "
                "Professional Subscribe via Stripe or pilot proforma",
                "You",
            ),
        ],
    )
    _bilingual_block(
        doc,
        "Realistic note: Steuernummer often arrives in days; USt-IdNr (VAT ID) can take 1–3 weeks. "
        "You can still register Gewerbe, open Stripe, and invoice at Intersolar with Steuerberater guidance "
        "(Kleinunternehmer §19 or regular 19% USt once Steuernummer is issued).",
        "Hinweis: Steuernummer oft in wenigen Tagen; USt-IdNr kann 1–3 Wochen dauern. "
        "Gewerbe-Anmeldung und Stripe sind trotzdem möglich — Abrechnung mit Steuerberater klären.",
    )

    # ── Gesellschaft question ──────────────────────────────────────────────────
    doc.add_heading("Do you need a Gesellschaft (company or association)? / Brauchen Sie eine Gesellschaft?", level=1)

    doc.add_heading("Short answer — no separate legal entity for Einzelunternehmen", level=2)
    _bilingual_block(
        doc,
        "For an Einzelunternehmen you do NOT need to found a UG, GmbH, or any other Kapitalgesellschaft. "
        "You operate in your own name (Mohammed Ismail Pasha), optionally with trading name “PVMath”. "
        "That is the simplest form — perfect for launching SaaS subscriptions before a future UG.",
        "Für ein Einzelunternehmen brauchen Sie KEINE UG, GmbH oder andere Kapitalgesellschaft. "
        "Sie treten in eigenem Namen auf (Mohammed Ismail Pasha), optional mit Marke „PVMath“. "
        "Das ist die schnellste Form für den Start von SaaS-Abos vor einer späteren UG.",
    )

    doc.add_heading("What you DO need (mandatory) / Was PFLICHT ist", level=2)
    _table(
        doc,
        ("Item", "English", "Deutsch", "Action"),
        [
            (
                "Gewerbeanmeldung",
                "Trade registration at Stadt Regensburg (Gewerbeamt)",
                "Anmeldung einer gewerblichen Tätigkeit",
                "Online or in person · often same day / next day",
            ),
            (
                "Finanzamt",
                "Tax registration questionnaire (Fragebogen zur steuerlichen Erfassung)",
                "Steuerliche Erfassung — Steuernummer + optional USt",
                "Auto-sent after Gewerbe or via ELSTER · Steuerberater helps",
            ),
            (
                "IHK",
                "Chamber of commerce membership",
                "Pflichtmitgliedschaft IHK Niederbayern",
                "Automatic when Gewerbe is registered — you receive a bill; "
                "you do NOT “found” an IHK",
            ),
            (
                "Trade register",
                "Handelsregister",
                "Handelsregister",
                "NOT required for standard Einzelunternehmen (only e.K. or GmbH/UG)",
            ),
        ],
    )

    doc.add_heading("Optional — NOT required for launch / Optional — NICHT Pflicht", level=2)
    _bullet(doc, "BSW-Solar, BVSE, or other Branchenverbände — useful later for networking, not for registration")
    _bullet(doc, "PVMath UG (haftungsbeschränkt) — defer until steady B2B revenue or EPC asks for HRB invoice")
    _bullet(doc, "Separate “business association” contract — no such entity needed for Einzelunternehmen")

    _p(doc, "Conclusion / Fazit:", bold=True)
    _bilingual_block(
        doc,
        "You do not need to draft or register a Gesellschaft. Register Gewerbe + tax, pay IHK when billed, "
        "optionally join BSW-Solar later for Intersolar networking.",
        "Sie müssen keine Gesellschaft gründen. Gewerbe + Steuer anmelden, IHK-Beitrag wenn abgerechnet, "
        "optional später Branchenverband für Networking.",
    )

    # ── Business activity ────────────────────────────────────────────────────
    doc.add_heading("Business activity (PVMath) — for Gewerbe & Finanzamt / Gewerbegegenstand", level=1)
    _p(
        doc,
        "Replace any generic template text with this PVMath-specific description. "
        "Your Steuerberater may shorten or add WZ codes.",
        bold=True,
    )

    doc.add_heading("English — business description", level=2)
    _p(
        doc,
        "PVMath operates a B2B software-as-a-service (SaaS) platform for utility-scale "
        "ground-mounted photovoltaic projects. The product brand is PVMath-Solar Site Intelligence Platform "
        "(tagline: “From site to system.”).",
    )
    _bullet(doc, "Target customers: solar EPCs, project developers, and engineering firms worldwide")
    _bullet(
        doc,
        "Scope: ground-mount only — fixed tilt, single-axis tracker, standard sites and Agri-PV "
        "(dual use). No rooftop, carport, floating, or BIPV.",
    )
    _bullet(
        doc,
        "Live modules: SiteIQ (rapid site screening — solar resource, terrain, flood, regulatory hints, "
        "capacity, PDF reports), TopoIQ (terrain slope and CAD export), YieldIQ (preliminary yield screening). "
        "Roadmap: RevenueIQ, LayoutIQ, ProcureIQ, FieldIQ.",
    )
    _bullet(
        doc,
        "Revenue model: freemium subscriptions (Free / Professional €149 per month / "
        "Developer €499 per month / Enterprise custom) via website pvmath.com and app siteiq.pvmath.com.",
    )
    _bullet(
        doc,
        "Technical activity: development and operation of web software; geospatial and engineering "
        "screening using open data (PVGIS, Copernicus DEM, OpenStreetMap); subscription billing and customer support.",
    )
    _bullet(
        doc,
        "Explicit limitation (important for liability): outputs are screening-grade for internal go/no-go "
        "decisions only — not bankable yield studies, not survey-grade terrain, not a substitute for "
        "LiDAR, PVsyst, or lender sign-off.",
    )
    _bullet(
        doc,
        "No field survey or drone operations in the current business phase (software and subscriptions only).",
    )

    doc.add_heading("Deutsch — Gewerbe-/Tätigkeitsbeschreibung (copy-paste)", level=2)
    _p(
        doc,
        "Entwicklung, Betrieb und Vertrieb von webbasierten Software-as-a-Service-Lösungen (SaaS) "
        "unter der Marke PVMath / SiteIQ zur technischen Vorprüfung und Standortanalyse von "
        "Freiflächen-Photovoltaik-Anlagen (Festachsanlagen, Nachführanlagen, Standard-Freifläche "
        "und Agri-PV).",
    )
    _p(
        doc,
        "Zielgruppe: EPC-Unternehmen, Projektentwickler und Ingenieurbüros im In- und Ausland.",
    )
    _p(
        doc,
        "Leistungsumfang: Bereitstellung und Wartung der Online-Module SiteIQ (Solar- und "
        "Standort-Screening inkl. PDF-Berichte), TopoIQ (Gelände-/Neigungsanalyse, CAD-Export) "
        "und YieldIQ (vorläufige Ertragsabschätzung); Abonnement-Vertrieb über pvmath.com und "
        "siteiq.pvmath.com; E-Mail-Support; Nutzung öffentlicher Geodaten (u. a. PVGIS, Copernicus).",
    )
    _p(
        doc,
        "Ausdrücklich keine bankfähigen Gutachten, keine vermessungsrechtlichen Leistungen, "
        "keine Drohnen- oder LiDAR-Feldaufnahmen in dieser Phase — ausschließlich Software und Abonnements.",
    )

    doc.add_heading("Suggested WZ / economic activity codes (confirm with Steuerberater)", level=2)
    _bullet(doc, "62.01 — Computer programming activities (Softwareentwicklung)")
    _bullet(doc, "62.09 — Other information technology service activities (SaaS-Betrieb)")
    _bullet(doc, "63.11 — Data processing, hosting (optional, if hosting emphasised)")
    _bullet(doc, "71.12 — Engineering-related consultancy (only if Steuerberater wants broader engineering wording)")

    doc.add_page_break()

    # ── Saturday checklist ───────────────────────────────────────────────────
    doc.add_heading("Saturday checklist — Einzelunternehmen / Checkliste Samstag", level=1)

    doc.add_heading("Before you leave the house / Vorbereitung", level=2)
    _bullet(doc, "Personalausweis or passport · Regensburg address proof if required")
    _bullet(doc, "Steuerberater call booked (30 min) — Kleinunternehmer §19 vs 19% USt decision")
    _bullet(doc, "Ideematec: written Nebentätigkeitsanzeige if contract requires it (see draft below)")
    _bullet(doc, "Bank account for business receipts (personal account OK if Steuerberater approved)")

    doc.add_heading("Step 1 — Gewerbeanmeldung Stadt Regensburg", level=2)
    _bilingual_block(
        doc,
        "Register online: Regensburg city portal “Gewerbe anmelden” or visit Bürgerservice. "
        "Legal form: Einzelunternehmen. Owner: Mohammed Ismail Pasha. "
        "Trading name (optional): PVMath. Paste business description from section above.",
        "Online über Stadt Regensburg „Gewerbe anmelden“ oder Bürgerservice. "
        "Rechtsform: Einzelunternehmen. Inhaber: Mohammed Ismail Pasha. "
        "Optional Fantasiename: PVMath. Gewerbebeschreibung aus diesem Dokument einfügen.",
    )
    _bullet(doc, "Start date: actual start or Monday 23 Jun if Steuerberater prefers")
    _bullet(doc, "Save confirmation PDF / receipt (Gewerbeschein comes after processing)")

    doc.add_heading("Step 2 — Finanzamt (steuerliche Erfassung)", level=2)
    _bilingual_block(
        doc,
        "After Gewerbe, Finanzamt Regensburg sends or offers ELSTER Fragebogen zur steuerlichen Erfassung. "
        "Complete with Steuerberater: expected revenue 2026, VAT option, activity description, IBAN.",
        "Nach Gewerbe-Anmeldung: Fragebogen zur steuerlichen Erfassung (ELSTER) — "
        "mit Steuerberater ausfüllen: Umsatzprognose, USt-Option, Tätigkeit, IBAN.",
    )
    _bullet(doc, "Request USt-IdNr if you invoice EU B2B customers (reverse charge)")
    _bullet(doc, "Steuernummer: needed for invoices — may arrive before USt-IdNr")

    doc.add_heading("Step 3 — Documents to update once Steuernummer known", level=2)
    _bullet(doc, "impressum.html — add operator name, address, Steuernummer when issued")
    _bullet(doc, "PVMath_Proforma_Invoice_Template.docx — fill Steuernummer / USt-IdNr fields")
    _bullet(doc, "Pilot agreement provider line: Mohammed Ismail Pasha, Einzelunternehmen, Regensburg")

    doc.add_page_break()

    # ── Sunday Stripe ────────────────────────────────────────────────────────
    doc.add_heading("Sunday — Stripe & payments / Stripe & Zahlungen", level=1)

    doc.add_heading("Stripe account setup", level=2)
    _table(
        doc,
        ("Field", "Value"),
        [
            ("Account type", "Individual / Einzelunternehmen"),
            ("Legal name", "Mohammed Ismail Pasha"),
            ("Trading name", "PVMath"),
            ("Business website", "https://pvmath.com"),
            ("Product description", "B2B SaaS solar site screening software"),
            ("Payout bank", "Your approved business/personal IBAN"),
            ("MCC", "Software / subscription services"),
        ],
    )

    doc.add_heading("Products & Payment Links", level=2)
    _bullet(doc, "Product: PVMath Professional — €149/month recurring (shared 75 analyses pool)")
    _bullet(doc, "Product: PVMath Developer — €499/month recurring (300 analyses, 5 seats)")
    _bullet(doc, "Create Stripe Payment Links → copy URLs for website and app")
    _bullet(doc, "Enable Customer Portal (cancel / invoice history)")
    _bullet(doc, "Enable SEPA Direct Debit for DE customers if Steuerberater agrees")

    doc.add_heading("Website & app wiring (dev task)", level=2)
    _bullet(doc, "index.html: Professional “Subscribe” → Stripe Payment Link (not #contact)")
    _bullet(doc, "index.html: Developer → Stripe link OR mailto with Developer subject until team sales process defined")
    _bullet(doc, "app Manage membership → same Stripe links + Customer Portal")
    _bullet(doc, "Phase 2 (same week): Supabase Edge Function webhook → auto profiles.plan on payment")

    doc.add_heading("Until webhook is live", level=2)
    _bilingual_block(
        doc,
        "Stripe sends payment email → you run SQL or use Manage membership team tools. "
        "Target: webhook before end of Intersolar week if possible.",
        "Stripe-Zahlungsmail → manuelle Freischaltung in Supabase bis Webhook live. "
        "Ziel: Webhook bis Ende der Messewoche.",
    )

    doc.add_page_break()

    # ── Intersolar marketing ───────────────────────────────────────────────────
    doc.add_heading("Tuesday — Intersolar marketing kit / Marketing Messe", level=1)

    doc.add_heading("Elevator pitch (30 seconds)", level=2)
    _p(
        doc,
        "EN: “PVMath SiteIQ screens utility-scale ground-mount sites in minutes — solar, terrain, "
        "and preliminary yield from open data. Fixed tilt, tracker, and Agri-PV. Free trial at siteiq.pvmath.com. "
        "Professional is €149/month for 75 pooled analyses across SiteIQ, TopoIQ, and YieldIQ.”",
    )
    _p(
        doc,
        "DE: „PVMath SiteIQ ist das Solar-Site-Screening für Freiflächen — Solar, Gelände und "
        "vorläufiger Ertrag in Minuten, auf Basis offener Daten. Festachse, Tracker, Agri-PV. "
        "Kostenlos testen auf siteiq.pvmath.com. Professional ab 149 €/Monat für 75 Analysen im gemeinsamen Kontingent.“",
    )

    doc.add_heading("Take to the fair", level=2)
    _bullet(doc, "QR code → https://siteiq.pvmath.com (test on phone before hall opens)")
    _bullet(doc, "Business cards: PVMath · Solar Site Intelligence · contact@pvmath.com · pvmath.com")
    _bullet(doc, "One-page PDF or phone demo: pin site in Bavaria or Spain — show PDF export")
    _bullet(doc, "LinkedIn posts scheduled Mon evening (founder + product — see marketing/ folder)")
    _bullet(doc, "Offer: Intersolar pilot — 3 months Professional with proforma or Stripe subscribe on spot")

    doc.add_heading("Lead capture", level=2)
    _bullet(doc, "Preferred: “Scan QR → free account → Subscribe in app or Stripe”")
    _bullet(doc, "Fallback: contact@pvmath.com with subject “Intersolar 2026 — Professional pilot”")
    _bullet(doc, "Note company + email on phone; activate within 24h if manual payment")

    doc.add_page_break()

    # ── Drafts ───────────────────────────────────────────────────────────────
    doc.add_heading("Draft — Nebentätigkeitsanzeige (if Ideematec requires)", level=1)
    _p(doc, "Send to HR / manager before or same day as Gewerbe registration.", bold=True)
    _p(
        doc,
        "Betreff: Anzeige einer nebenberuflichen selbständigen Tätigkeit\n\n"
        "Sehr geehrte Damen und Herren,\n\n"
        "hiermit zeige ich gemäß meinem Arbeitsvertrag an, dass ich eine nebenberufliche "
        "selbständige Tätigkeit als Einzelunternehmer aufnehme.\n\n"
        "Tätigkeit: Entwicklung und Vertrieb einer B2B-Software (PVMath / SiteIQ) zur "
        "technischen Vorprüfung von Freiflächen-Photovoltaik-Projekten — reine Software-Abonnements, "
        "keine konkurrierende EPC-Leistung, keine Tätigkeit in Arbeitszeit oder mit Firmenressourcen.\n\n"
        "Marken/Domains: pvmath.com, siteiq.pvmath.com (privat finanziert).\n\n"
        "Ich bitte um schriftliche Kenntnisnahme. Bei Rückfragen stehe ich gerne zur Verfügung.\n\n"
        "Mit freundlichen Grüßen\n"
        "Mohammed Ismail Pasha",
    )

    doc.add_heading("Draft — Email to Steuerberater (Saturday morning)", level=1)
    _p(
        doc,
        "Subject: PVMath Einzelunternehmen — Gewerbe heute · USt · Intersolar Dienstag\n\n"
        "Hallo [Name],\n\n"
        "ich melde heute mein Einzelunternehmen (PVMath, SaaS Solar-Screening, siehe Tätigkeitsbeschreibung) "
        "in Regensburg an. Bitte kurz bestätigen:\n"
        "1) Kleinunternehmer §19 vs reguläre USt 19% für 2026\n"
        "2) USt-IdNr sofort beantragen ja/nein\n"
        "3) Rechnungsstellung ab Dienstag (Intersolar) mit Steuernummer wenn USt-IdNr noch fehlt\n"
        "4) Stripe-Auszahlungen auf [IBAN] — buchhalterisch ok?\n\n"
        "Vielen Dank,\n"
        "Ismail",
    )

    doc.add_heading("Draft — Impressum operator line (after registration)", level=1)
    _p(
        doc,
        "EN/DE: Mohammed Ismail Pasha · Einzelunternehmen · [Straße, PLZ Regensburg] · "
        "Germany · contact@pvmath.com · USt-IdNr.: [DE…] · Steuernummer: […]",
    )

    doc.add_heading("Checklist summary — print and tick", level=1)
    items = [
        "☐ Steuerberater briefed (VAT path)",
        "☐ Ideematec Nebentätigkeit disclosed if required",
        "☐ Gewerbe registered — confirmation saved",
        "☐ Fragebogen steuerliche Erfassung submitted",
        "☐ Stripe account verified",
        "☐ Payment Links created (Pro + Dev)",
        "☐ Website Subscribe button → Stripe",
        "☐ App Manage membership → Stripe",
        "☐ Impressum updated",
        "☐ QR + demo project tested",
        "☐ Intersolar pitch rehearsed (EN + DE)",
    ]
    for item in items:
        _bullet(doc, item)

    _p(doc, "Document generated for PVMath internal use · June 2026", center=True)
    return doc


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    build().save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
